import json
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from datetime import datetime
import os

class ReportGenerator:
    def __init__(self, agency_name="Lear Cyber Tech", logo_path=None):
        self.agency_name = agency_name
        self.logo_path = logo_path or "/home/ubuntu/upload/CyberLear-Logo.png"
        self.brand_color = "#1e3a8a"  # Lear Cyber Tech blue
        
    def generate_excel_report(self, data, output_path, template_name):
        """Generate branded Excel report"""
        wb = Workbook()
        ws = wb.active
        ws.title = "Report"
        
        # Header styling
        header_font = Font(bold=True, color="FFFFFF", size=14)
        header_fill = PatternFill(start_color="1e3a8a", end_color="1e3a8a", fill_type="solid")
        
        # Title
        ws.merge_cells('A1:F1')
        ws['A1'] = f"{self.agency_name} - {template_name}"
        ws['A1'].font = header_font
        ws['A1'].fill = header_fill
        ws['A1'].alignment = Alignment(horizontal='center')
        
        # Client information
        row = 3
        ws[f'A{row}'] = "Client Information"
        ws[f'A{row}'].font = Font(bold=True, size=12)
        row += 1
        
        client_info = data.get('client_info', {})
        for key, value in client_info.items():
            ws[f'A{row}'] = key.replace('_', ' ').title()
            ws[f'B{row}'] = value
            row += 1
        
        row += 1
        
        # Assessment data
        ws[f'A{row}'] = "Assessment Results"
        ws[f'A{row}'].font = Font(bold=True, size=12)
        row += 1
        
        answers = data.get('answers', {})
        for question_id, answer in answers.items():
            ws[f'A{row}'] = f"Question {question_id}"
            ws[f'B{row}'] = str(answer)
            row += 1
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        wb.save(output_path)
        return output_path
    
    def generate_word_report(self, data, output_path, template_name):
        """Generate branded Word report"""
        doc = Document()
        
        # Header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = f"{self.agency_name} - {template_name}"
        
        # Title
        title = doc.add_heading(f"{template_name} Report", 0)
        title.alignment = 1  # Center alignment
        
        # Generated date
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph("")
        
        # Client information
        doc.add_heading("Client Information", level=1)
        client_info = data.get('client_info', {})
        for key, value in client_info.items():
            doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
        
        doc.add_paragraph("")
        
        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(
            "This report provides a comprehensive assessment based on the completed template. "
            "The following sections detail the findings and recommendations."
        )
        
        doc.add_paragraph("")
        
        # Assessment Results
        doc.add_heading("Assessment Results", level=1)
        answers = data.get('answers', {})
        for question_id, answer in answers.items():
            doc.add_paragraph(f"Question {question_id}: {answer}")
        
        # Recommendations
        doc.add_heading("Recommendations", level=1)
        doc.add_paragraph(
            "Based on the assessment results, we recommend the following actions to improve "
            "your cybersecurity posture and compliance status."
        )
        
        # Footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"© {datetime.now().year} {self.agency_name}. All rights reserved."
        
        doc.save(output_path)
        return output_path
    
    def generate_pdf_report(self, data, output_path, template_name):
        """Generate branded PDF report"""
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # Center
            textColor=colors.HexColor(self.brand_color)
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            textColor=colors.HexColor(self.brand_color)
        )
        
        # Title
        story.append(Paragraph(f"{template_name} Report", title_style))
        story.append(Paragraph(f"Prepared by {self.agency_name}", styles['Normal']))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Client Information
        story.append(Paragraph("Client Information", heading_style))
        client_info = data.get('client_info', {})
        
        client_data = []
        for key, value in client_info.items():
            client_data.append([key.replace('_', ' ').title(), str(value)])
        
        if client_data:
            client_table = Table(client_data, colWidths=[2*inch, 4*inch])
            client_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.brand_color)),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(client_table)
        
        story.append(Spacer(1, 20))
        
        # Assessment Results
        story.append(Paragraph("Assessment Results", heading_style))
        answers = data.get('answers', {})
        
        assessment_data = []
        for question_id, answer in answers.items():
            assessment_data.append([f"Question {question_id}", str(answer)])
        
        if assessment_data:
            assessment_table = Table(assessment_data, colWidths=[2*inch, 4*inch])
            assessment_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.brand_color)),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(assessment_table)
        
        story.append(Spacer(1, 20))
        
        # Recommendations
        story.append(Paragraph("Recommendations", heading_style))
        story.append(Paragraph(
            "Based on the assessment results, we recommend implementing the following "
            "security measures to enhance your cybersecurity posture:",
            styles['Normal']
        ))
        
        # Footer
        story.append(Spacer(1, 40))
        story.append(Paragraph(
            f"© {datetime.now().year} {self.agency_name}. All rights reserved.",
            styles['Normal']
        ))
        
        doc.build(story)
        return output_path

class TemplateProcessor:
    def __init__(self):
        self.report_generator = ReportGenerator()
    
    def create_blank_template(self, template_type, questions, output_path):
        """Create blank template for download"""
        if template_type == "excel":
            return self._create_excel_template(questions, output_path)
        elif template_type == "word":
            return self._create_word_template(questions, output_path)
        else:
            return self._create_pdf_template(questions, output_path)
    
    def _create_excel_template(self, questions, output_path):
        """Create blank Excel template"""
        wb = Workbook()
        ws = wb.active
        ws.title = "Template"
        
        # Header
        ws['A1'] = "Lear Cyber Tech - Assessment Template"
        ws['A1'].font = Font(bold=True, size=14)
        
        # Instructions
        ws['A3'] = "Instructions:"
        ws['A4'] = "1. Fill in your information in the Client Information section"
        ws['A5'] = "2. Answer all questions in the Assessment section"
        ws['A6'] = "3. Upload the completed template to generate your report"
        
        # Client Information section
        row = 8
        ws[f'A{row}'] = "CLIENT INFORMATION"
        ws[f'A{row}'].font = Font(bold=True, size=12)
        row += 1
        
        client_fields = [
            "Company Name",
            "Contact Person",
            "Email Address",
            "Phone Number",
            "Industry",
            "Assessment Date"
        ]
        
        for field in client_fields:
            ws[f'A{row}'] = field
            ws[f'B{row}'] = "[Enter your information here]"
            row += 1
        
        row += 2
        
        # Questions section
        ws[f'A{row}'] = "ASSESSMENT QUESTIONS"
        ws[f'A{row}'].font = Font(bold=True, size=12)
        row += 1
        
        for i, question in enumerate(questions, 1):
            ws[f'A{row}'] = f"Q{i}: {question.get('text', '')}"
            ws[f'B{row}'] = "[Your answer here]"
            if question.get('type') == 'multiple_choice':
                options = question.get('options', [])
                ws[f'C{row}'] = f"Options: {', '.join(options)}"
            row += 1
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 80)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        wb.save(output_path)
        return output_path
    
    def _create_word_template(self, questions, output_path):
        """Create blank Word template"""
        doc = Document()
        
        # Title
        title = doc.add_heading("Lear Cyber Tech - Assessment Template", 0)
        title.alignment = 1
        
        # Instructions
        doc.add_heading("Instructions", level=1)
        instructions = [
            "1. Fill in your information in the Client Information section",
            "2. Answer all questions in the Assessment section",
            "3. Upload the completed template to generate your report"
        ]
        for instruction in instructions:
            doc.add_paragraph(instruction)
        
        # Client Information
        doc.add_heading("Client Information", level=1)
        client_fields = [
            "Company Name: ________________________________",
            "Contact Person: ________________________________",
            "Email Address: ________________________________",
            "Phone Number: ________________________________",
            "Industry: ________________________________",
            "Assessment Date: ________________________________"
        ]
        for field in client_fields:
            doc.add_paragraph(field)
        
        # Questions
        doc.add_heading("Assessment Questions", level=1)
        for i, question in enumerate(questions, 1):
            doc.add_paragraph(f"Q{i}: {question.get('text', '')}")
            if question.get('type') == 'multiple_choice':
                options = question.get('options', [])
                for option in options:
                    doc.add_paragraph(f"☐ {option}")
            else:
                doc.add_paragraph("Answer: ________________________________")
            doc.add_paragraph("")
        
        doc.save(output_path)
        return output_path
    
    def process_filled_template(self, file_path, file_type):
        """Process filled template and extract data"""
        if file_type == "xlsx":
            return self._process_excel_template(file_path)
        elif file_type == "docx":
            return self._process_word_template(file_path)
        else:
            return {"error": "Unsupported file type"}
    
    def _process_excel_template(self, file_path):
        """Process filled Excel template"""
        try:
            from openpyxl import load_workbook
            wb = load_workbook(file_path)
            ws = wb.active
            
            extracted_data = {
                "client_info": {},
                "answers": {},
                "metadata": {
                    "processed_at": datetime.now().isoformat(),
                    "file_type": "excel"
                }
            }
            
            # Extract client information (assuming it starts at row 9)
            client_fields = [
                "company_name", "contact_person", "email_address",
                "phone_number", "industry", "assessment_date"
            ]
            
            for i, field in enumerate(client_fields, 9):
                value = ws[f'B{i}'].value
                if value and not str(value).startswith("["):
                    extracted_data["client_info"][field] = str(value)
            
            # Extract answers (find questions section)
            for row in range(1, ws.max_row + 1):
                cell_value = ws[f'A{row}'].value
                if cell_value and "ASSESSMENT QUESTIONS" in str(cell_value):
                    # Start extracting answers from next row
                    answer_row = row + 1
                    question_num = 1
                    while answer_row <= ws.max_row:
                        question_cell = ws[f'A{answer_row}'].value
                        answer_cell = ws[f'B{answer_row}'].value
                        
                        if question_cell and str(question_cell).startswith(f"Q{question_num}"):
                            if answer_cell and not str(answer_cell).startswith("["):
                                extracted_data["answers"][f"q{question_num}"] = str(answer_cell)
                            question_num += 1
                        
                        answer_row += 1
                        if answer_row > row + 50:  # Safety limit
                            break
                    break
            
            return extracted_data
            
        except Exception as e:
            return {"error": f"Failed to process Excel file: {str(e)}"}
    
    def _process_word_template(self, file_path):
        """Process filled Word template"""
        try:
            from docx import Document
            doc = Document(file_path)
            
            extracted_data = {
                "client_info": {},
                "answers": {},
                "metadata": {
                    "processed_at": datetime.now().isoformat(),
                    "file_type": "word"
                }
            }
            
            # Extract text from all paragraphs
            all_text = []
            for paragraph in doc.paragraphs:
                all_text.append(paragraph.text)
            
            # Simple extraction logic (can be enhanced)
            for i, text in enumerate(all_text):
                if "Company Name:" in text:
                    extracted_data["client_info"]["company_name"] = text.split(":")[-1].strip()
                elif "Contact Person:" in text:
                    extracted_data["client_info"]["contact_person"] = text.split(":")[-1].strip()
                elif "Email Address:" in text:
                    extracted_data["client_info"]["email_address"] = text.split(":")[-1].strip()
                elif text.startswith("Q") and ":" in text:
                    # Extract question and look for answer in next paragraphs
                    question_num = text.split(":")[0].replace("Q", "")
                    if i + 1 < len(all_text) and "Answer:" in all_text[i + 1]:
                        answer = all_text[i + 1].split(":")[-1].strip()
                        if answer and not answer.startswith("_"):
                            extracted_data["answers"][f"q{question_num}"] = answer
            
            return extracted_data
            
        except Exception as e:
            return {"error": f"Failed to process Word file: {str(e)}"}

# Initialize global instances
template_processor = TemplateProcessor()
report_generator = ReportGenerator()

